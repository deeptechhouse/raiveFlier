"""Unit tests for FormatConverter service.

Tests format-specific conversion methods using temporary files and mocking
external CLI tools (Calibre, djvutxt, LibreOffice) to avoid system deps.
"""

from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest

from tools.raive_feeder.services.format_converter import FormatConverter


@pytest.fixture
def converter():
    return FormatConverter()


class TestConvertDispatch:
    """Tests for the main convert() dispatcher."""

    @pytest.mark.asyncio
    async def test_unsupported_format_raises(self, converter):
        """Converting an unsupported format should raise ValueError."""
        with pytest.raises(ValueError, match="No converter available"):
            await converter.convert("/tmp/test.xyz", ".xyz")

    @pytest.mark.asyncio
    async def test_dispatch_routes_to_rtf(self, converter):
        """RTF files should be dispatched to _convert_rtf."""
        with patch.object(converter, "_convert_rtf", new_callable=AsyncMock) as mock:
            mock.return_value = ("/tmp/out.txt", "txt")
            result = await converter.convert("/tmp/test.rtf", ".rtf")
            mock.assert_called_once_with("/tmp/test.rtf")
            assert result[1] == "txt"


class TestConvertDocx:
    """Tests for DOCX conversion (pure Python, no CLI needed)."""

    @pytest.mark.asyncio
    async def test_convert_docx(self, converter):
        """DOCX with paragraphs should produce non-empty text output."""
        # Create a minimal DOCX file using python-docx.
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Test paragraph one.")
        doc.add_paragraph("Test paragraph two.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        output_path, fmt = await converter._convert_docx(tmp_path)

        assert fmt == "txt"
        assert Path(output_path).exists()
        content = Path(output_path).read_text()
        assert "Test paragraph one" in content
        assert "Test paragraph two" in content


class TestConvertRtf:
    """Tests for RTF conversion (pure Python via striprtf)."""

    @pytest.mark.asyncio
    async def test_convert_rtf(self, converter):
        """RTF content should produce plain text output."""
        try:
            from striprtf.striprtf import rtf_to_text  # noqa: F401
        except ImportError:
            pytest.skip("striprtf not installed")

        # Minimal RTF content.
        rtf_content = r"{\rtf1\ansi Test RTF content here.}"
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".rtf", delete=False, encoding="utf-8"
        ) as tmp:
            tmp.write(rtf_content)
            tmp_path = tmp.name

        output_path, fmt = await converter._convert_rtf(tmp_path)

        assert fmt == "txt"
        assert Path(output_path).exists()
        content = Path(output_path).read_text()
        assert "Test RTF content" in content


class TestConvertMobi:
    """Tests for MOBI conversion (requires Calibre CLI)."""

    @pytest.mark.asyncio
    async def test_mobi_without_calibre_raises(self, converter):
        """MOBI conversion without Calibre installed should raise RuntimeError."""
        with patch("shutil.which", return_value=None):
            with pytest.raises(RuntimeError, match="Calibre not installed"):
                await converter._convert_mobi("/tmp/test.mobi")


class TestConvertDjvu:
    """Tests for DJVU conversion (requires djvulibre CLI)."""

    @pytest.mark.asyncio
    async def test_djvu_without_djvulibre_raises(self, converter):
        """DJVU conversion without djvulibre installed should raise RuntimeError."""
        with patch("shutil.which", return_value=None):
            with pytest.raises(RuntimeError, match="djvulibre not installed"):
                await converter._convert_djvu("/tmp/test.djvu")


class TestConvertDoc:
    """Tests for DOC conversion (requires LibreOffice CLI)."""

    @pytest.mark.asyncio
    async def test_doc_without_libreoffice_raises(self, converter):
        """DOC conversion without LibreOffice should raise RuntimeError."""
        with patch("shutil.which", return_value=None):
            with pytest.raises(RuntimeError, match="LibreOffice not installed"):
                await converter._convert_doc("/tmp/test.doc")

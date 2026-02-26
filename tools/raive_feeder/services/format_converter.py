"""Format conversion service for unsupported document types.

# ─── DESIGN ────────────────────────────────────────────────────────────
#
# FormatConverter bridges the gap between file formats that users upload
# (MOBI, DJVU, DOC, DOCX, RTF) and the formats that raiveFlier's
# IngestionService natively handles (TXT, PDF, EPUB).
#
# Conversion strategy per format:
#   MOBI  → EPUB   via Calibre's ebook-convert CLI
#   DJVU  → TXT    via djvutxt CLI (from djvulibre)
#   DOCX  → TXT    via python-docx library (pure Python)
#   DOC   → TXT    via LibreOffice CLI (headless)
#   RTF   → TXT    via striprtf library (pure Python)
#
# External tool detection happens at call time, not at init, so the
# service degrades gracefully — if Calibre isn't installed, MOBI
# conversion fails with a clear error while other formats still work.
#
# Pattern: Strategy (format → converter function dispatch).
# ──────────────────────────────────────────────────────────────────────
"""

from __future__ import annotations

import asyncio
import shutil
import tempfile
from pathlib import Path

import structlog

logger = structlog.get_logger(logger_name=__name__)


class FormatConverter:
    """Converts unsupported document formats to formats the ingestion pipeline handles.

    All conversion methods are async to avoid blocking the event loop
    when shelling out to CLI tools (Calibre, djvutxt, LibreOffice).
    """

    async def convert(self, file_path: str, suffix: str) -> tuple[str, str]:
        """Convert a file to a processable format.

        Parameters
        ----------
        file_path:
            Path to the uploaded file.
        suffix:
            Lowercase file extension including dot (e.g. ".mobi").

        Returns
        -------
        tuple[str, str]
            (converted_file_path, output_format) where output_format is
            "epub" or "txt".
        """
        converters = {
            ".mobi": self._convert_mobi,
            ".djvu": self._convert_djvu,
            ".docx": self._convert_docx,
            ".doc": self._convert_doc,
            ".rtf": self._convert_rtf,
        }

        converter = converters.get(suffix)
        if converter is None:
            raise ValueError(f"No converter available for format: {suffix}")

        return await converter(file_path)

    async def _convert_mobi(self, file_path: str) -> tuple[str, str]:
        """MOBI → EPUB via Calibre's ebook-convert CLI.

        Calibre's ebook-convert is the most reliable MOBI decoder.
        It handles DRM-free MOBI/AZW3 files and produces clean EPUB.
        """
        if not shutil.which("ebook-convert"):
            raise RuntimeError(
                "Calibre not installed. Install via: brew install calibre "
                "(macOS) or apt install calibre (Linux)"
            )

        output_path = tempfile.mktemp(suffix=".epub")
        proc = await asyncio.create_subprocess_exec(
            "ebook-convert", file_path, output_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        _, stderr = await proc.communicate()

        if proc.returncode != 0:
            raise RuntimeError(f"ebook-convert failed: {stderr.decode()[:500]}")

        logger.info("mobi_converted", output=output_path)
        return output_path, "epub"

    async def _convert_djvu(self, file_path: str) -> tuple[str, str]:
        """DJVU → plain text via djvutxt CLI (from djvulibre).

        djvutxt extracts the hidden text layer from DJVU files.
        If the DJVU has no OCR text layer, the output will be empty.
        """
        if not shutil.which("djvutxt"):
            raise RuntimeError(
                "djvulibre not installed. Install via: brew install djvulibre "
                "(macOS) or apt install djvulibre-bin (Linux)"
            )

        output_path = tempfile.mktemp(suffix=".txt")
        proc = await asyncio.create_subprocess_exec(
            "djvutxt", file_path, output_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        _, stderr = await proc.communicate()

        if proc.returncode != 0:
            raise RuntimeError(f"djvutxt failed: {stderr.decode()[:500]}")

        logger.info("djvu_converted", output=output_path)
        return output_path, "txt"

    async def _convert_docx(self, file_path: str) -> tuple[str, str]:
        """DOCX → plain text via python-docx (pure Python, no CLI needed).

        python-docx reads the XML inside the DOCX zip archive and
        extracts paragraph text.  Formatting is stripped.
        """
        from docx import Document

        doc = Document(file_path)
        text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())

        output_path = tempfile.mktemp(suffix=".txt")
        Path(output_path).write_text(text, encoding="utf-8")

        logger.info("docx_converted", paragraphs=len(doc.paragraphs), output=output_path)
        return output_path, "txt"

    async def _convert_doc(self, file_path: str) -> tuple[str, str]:
        """DOC → plain text via LibreOffice CLI (headless mode).

        Legacy .doc files use a proprietary binary format.  LibreOffice
        is the most reliable open-source tool for reading them.
        """
        lo_cmd = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo_cmd:
            raise RuntimeError(
                "LibreOffice not installed. Install via: brew install --cask libreoffice "
                "(macOS) or apt install libreoffice (Linux)"
            )

        output_dir = tempfile.mkdtemp()
        proc = await asyncio.create_subprocess_exec(
            lo_cmd, "--headless", "--convert-to", "txt",
            "--outdir", output_dir, file_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        await proc.communicate()

        # LibreOffice names the output file with .txt extension.
        stem = Path(file_path).stem
        output_path = str(Path(output_dir) / f"{stem}.txt")

        if not Path(output_path).exists():
            raise RuntimeError("LibreOffice conversion produced no output")

        logger.info("doc_converted", output=output_path)
        return output_path, "txt"

    async def _convert_rtf(self, file_path: str) -> tuple[str, str]:
        """RTF → plain text via striprtf library (pure Python).

        striprtf parses RTF control words and extracts the plain text
        content without requiring any external dependencies.
        """
        from striprtf.striprtf import rtf_to_text

        rtf_content = Path(file_path).read_text(encoding="utf-8", errors="replace")
        text = rtf_to_text(rtf_content)

        output_path = tempfile.mktemp(suffix=".txt")
        Path(output_path).write_text(text, encoding="utf-8")

        logger.info("rtf_converted", output=output_path)
        return output_path, "txt"
